from copy import deepcopy
from pathlib import Path

from docx import Document


path = Path("Устав проекта AI Regulation Monitor.docx")
document = Document(path)
team_table = document.tables[2]

if any("Тестировщик" in cell.text for row in team_table.rows for cell in row.cells):
    raise RuntimeError("Строка тестировщика уже присутствует в таблице")

template_row = team_table.rows[-1]
new_row_xml = deepcopy(template_row._tr)
team_table._tbl.append(new_row_xml)
new_row = team_table.rows[-1]

values = [
    "Тестировщик (QA Engineer)",
    (
        "Проверка функциональности, интерфейса и адаптивности; тестирование "
        "авторизации, API, безопасности и производительности; проверка "
        "качества и юридической релевантности результатов; регистрация "
        "дефектов и контроль их исправления."
    ),
    (
        "Ручное и автоматизированное тестирование; тестирование Flask и API; "
        "UX/UI; основы информационной безопасности; нагрузочное тестирование; "
        "подготовка тест-кейсов и отчетов."
    ),
    "Тестировщик программного обеспечения",
    (
        "Обеспечивает качество, надежность и безопасность приложения, "
        "проверяет корректность отбора материалов и готовность проекта "
        "к эксплуатации."
    ),
]

for cell, value in zip(new_row.cells, values):
    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = value
    else:
        paragraph.add_run(value)
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)

document.save(path)
print(f"Updated: {path}")
