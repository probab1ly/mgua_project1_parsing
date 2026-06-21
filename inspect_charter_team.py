from pathlib import Path
from docx import Document


path = Path("Устав проекта AI Regulation Monitor.docx")
document = Document(path)

print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)}")
for table_index, table in enumerate(document.tables):
    rows = []
    for row_index, row in enumerate(table.rows):
        text = " | ".join(
            " ".join(cell.text.split())
            for cell in row.cells
        )
        rows.append((row_index, text))
    combined = "\n".join(text for _, text in rows).lower()
    if any(word in combined for word in ("владелец", "разработчик", "команда", "роль")):
        print(f"\nTABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, text in rows:
            print(f"{row_index}: {text}")
